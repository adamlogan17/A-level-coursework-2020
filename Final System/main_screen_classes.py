import sys

import tkinter as tk

from tkinter import messagebox

import tkinter.ttk as ttk

from permission_classes import *

from tkinter import filedialog

from datetime import datetime, timedelta, time

from PIL import Image, ImageTk

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import matplotlib.ticker as ticker

def get_file_location():
    '''This will ask the user to select a directory and call the function 'backup' '''
    
    file = "Hotel_management_system.db"
    drctory_of_backup = filedialog.askdirectory(title = "Select the folder you would like the backup to save")
    today = datetime.today() # stores todays date and time the script was run.
    current_date = today.strftime("%d-%m-%Y") # the date is in the format 'd-m-y' and not 'd/m/y' because windows will not allow file names to contian the character '/'.
    backup(current_date,file,drctory_of_backup)

def set_image(frame,img_name, width, height):
    '''Places the image within the frame with a specified width and height'''
    clear_frame_or_window(frame)
    if img_name != None:
        img_name = Image.open(img_name)

        img_name = img_name.resize((width, height), Image.ANTIALIAS)
        
        img = ImageTk.PhotoImage(img_name) 

        photo = tk.Label(frame,image=img)

        photo.image = img # needs to keep a reference for the image so it is not discarded. see more at "http://effbot.org/pyfaq/why-do-my-tkinter-images-not-appear.htm"

        photo['bg'] = photo.master['bg']

        photo.pack(side=tk.BOTTOM)
    else:
        text = """

Sorry, there are no bookings
at this time


                """
        text = tk.Message(frame,text=text,background="#d9d9d9",aspect=1000,
                                  font="-family Arial -size 30", foreground = "#000000",anchor=tk.CENTER)
        text.pack(expand=True,fill = tk.BOTH)

def get_data_for_gantt(month, year):
    '''Calculates the data for the 'Tape Chart' and returns an image name'''
    try:
        year = int(year)
    except ValueError:
        messagebox.showerror("Error","Please enter a year, not letter or special characters.")
        return None
    with sqlite3.connect("Hotel_Management_system.db") as db:
        cursor=db.cursor()
        cursor.execute("select RoomID,Room_Number from Room")
        rooms=cursor.fetchall()
        db.commit()

    room_info = {}
    for i in range(len(rooms)):
        with sqlite3.connect("Hotel_Management_system.db") as db:
            cursor=db.cursor()
            cursor.execute("select Check_In,Check_Out from Booking where RoomID = {0}".format(rooms[i][0]))
            dates=cursor.fetchall()
            db.commit()
        room_info[rooms[i][1]] = dates
    
    check_days = []
    today = datetime.today()
    count = 0
    for key, value in room_info.items():
        orig_value = value[:]
        for i in range(len(value)):
            start_date = None
            check_in = datetime.strptime(orig_value[i][0],"%d/%m/%Y")
            check_out = datetime.strptime(orig_value[i][1],"%d/%m/%Y")
    
            if (check_in.month < month < check_out.month and check_in.year == year) or (check_in.year < year < check_out.year) or (check_out.month > month and check_in.year < year and check_out.year == year) or (check_in.month < month and check_out.year > year and check_in.year == year):
                start_date = 1
                difference = 31
            elif check_in.month == month and check_in.month == check_out.month and check_in.year == year and check_in.year == check_out.year:
                start_date = check_in.day
                difference = int(check_out.day) - int(check_in.day)
            elif (check_in.month == month and (check_in.month != check_out.month or check_out.year != year) and check_in.year == year):
                start_date = check_in.day
                difference = 31 - int(check_in.day)
            elif check_out.month == month and check_out.year == year:
                start_date = 0
                difference = check_out.day
            else:
                value.remove(orig_value[i])
                
            if start_date != None:
                reference = value[value.index(orig_value[i])][ : 0] + (start_date,) + value[value.index(orig_value[i])][0 + 1 : ]
                value[value.index(orig_value[i])] = value[value.index(orig_value[i])][ : 0] + (start_date,) + value[value.index(orig_value[i])][0 + 1 : ]
                value[value.index(reference)] = value[value.index(reference)][ : 1] + (difference,) + value[value.index(reference)][1 + 1 : ]
                
        if value == []:
            count = count+1
            room_info[key] = [(0,0)]

    if count != len(room_info):
        ## Below sets the correct number of days for the gantt chart.  
        if month in [4,6,11,9]:
            days_in_month = (1,30)
        elif month in [1,3,5,7,8,10,12]:
            days_in_month = (1,31)
        elif month == 2 and year%4 == 0:
            days_in_month = (1,29)
        else:
            days_in_month = (1,28)

        gantt_chart_img_name = gantt_chart(room_info, days_in_month)

        return gantt_chart_img_name


def retrn_table_names(db_name): # Needs to be a string. db_name needs to be the directory of the database including the file extension '.db'. This will only work if it is a db.
    '''This will return all the table names of a database as a list.'''
    lst_of_tables=[]
    with sqlite3.connect(db_name) as db:
        cursor = db.cursor()
        cursor.execute("select name from sqlite_master")
        lst_of_tables_in_tuple=cursor.fetchall()
        for table_name in lst_of_tables_in_tuple: # iterates through the tuple to only get the names of the tables and nothing else.
            lst_of_tables.append(table_name[0])
        return lst_of_tables # returns the names of the tables in a list.

def quit_command():
    '''This will display a quit message box and if
       the user clicks yes then it will destroy the window.'''
    if messagebox.askyesno("Quit","Are you sure you want to QUIT?"):
        os._exit(0) ## This stops the script from running. It also stops the shell form asking the user to 'kill' the program. 'https://stackoverflow.com/questions/19782075/how-to-stop-terminate-a-python-script-from-running/34029481'


class MainMenu:
    def __init__(self, top=None, message=None, users_record=None):
        '''This class populates the template of all the other classes.'''
        
        _bgcolor = '#d9d9d9'  # X11 color: 'gray85'
        _fgcolor = '#000000'  # X11 color: 'black'
        _compcolor = '#d9d9d9' # X11 color: 'gray85'
        _ana1color = '#d9d9d9' # X11 color: 'gray85'
        _ana2color = '#ececec' # Closest X11 color: 'gray92'
        font14 = "-family Arial -size 10 -weight bold -slant roman "  \
            "-underline 1 -overstrike 0"
        font1 = "-family Arial -size 8 -weight normal -slant roman "  \
            "-underline 0 -overstrike 0"
        self.style = ttk.Style()
        if sys.platform == "win32":
            self.style.theme_use('winnative')
        self.style.configure('.',background=_bgcolor)
        self.style.configure('.',foreground=_fgcolor)
        self.style.configure('.',font="TkDefaultFont")
        self.style.map('.',background=
            [('selected', _compcolor), ('active',_ana2color)])

        #top.attributes("-fullscreen", True) # This expands the window to full screen but there is no bar at the top to minimise it. check out 'http://www.tcl.tk/man/tcl/TkCmd/wm.htm#m8'
        #top.resizable(False, False) # Does not allow the window to re-size. Got from 'https://stackoverflow.com/questions/37446710/how-to-make-a-tkinter-window-not-resizable'
        top.geometry("684x450+283+165")
        top.title("Hotel Management System")
        top.configure(background="#d9d9d9")
        top.configure(highlightbackground="#d9d9d9")
        top.configure(highlightcolor="black")

        self.TSeparator1 = ttk.Separator(top)
        self.TSeparator1.place(relx=0.132, rely=-0.044, relheight=1.711)
        self.TSeparator1.configure(orient="vertical")

        self.TSeparator2 = ttk.Separator(top)
        self.TSeparator2.place(relx=-0.058, rely=0.244, relwidth=2.047)

        self.ImageTFrame1 = tk.Frame(top)
        self.ImageTFrame1.place(relx=0.000, rely=0.000, relheight=0.243
                , relwidth=0.129)
        #self.ImageTFrame1.configure(relief='groove')
        self.ImageTFrame1.configure(borderwidth="2")
        self.ImageTFrame1.configure(background="#d9d9d9")
        #self.ImageTFrame1.configure(relief="groove")
        
        img_name = Image.open("hotel_logo.png")
        img_name = img_name.resize((80, 100), Image.ANTIALIAS)
        img = ImageTk.PhotoImage(img_name)
        photo = tk.Label(self.ImageTFrame1,image=img)
        photo.configure(background="#d9d9d9")
        photo.image = img # needs to keep a reference for the image so it is not discarded. see more at "http://effbot.org/pyfaq/why-do-my-tkinter-images-not-appear.htm"
        #photo['bg'] = photo.master['bg']
        photo.pack(expand=True,fill="both")
        
        self.menubar = tk.Menu(top,font="TkMenuFont",bg=_bgcolor,fg=_fgcolor)
        top.configure(menu = self.menubar)

        self.Quit_Button = tk.Button(top)
        self.Quit_Button.place(relx=0.015, rely=0.911, height=31, width=71)
        self.Quit_Button.configure(activebackground="#ececec")
        self.Quit_Button.configure(activeforeground="#000000")
        self.Quit_Button.configure(background="#d9d9d9")
        self.Quit_Button.configure(disabledforeground="#a3a3a3")
        self.Quit_Button.configure(font=font14)
        self.Quit_Button.configure(foreground="#000000")
        self.Quit_Button.configure(highlightbackground="#d9d9d9")
        self.Quit_Button.configure(highlightcolor="black")
        self.Quit_Button.configure(pady="0")
        self.Quit_Button.configure(text='''Quit''')
        self.Quit_Button.configure(command=quit_command)
        self.Quit_Button.configure(cursor='hand2')

        self.Welcome_Label = tk.Label(top)
        self.Welcome_Label.place(relx=0.146, rely=0.022, height=100, width=550)
        self.Welcome_Label.configure(activebackground="#f9f9f9")
        self.Welcome_Label.configure(activeforeground="black")
        self.Welcome_Label.configure(background="#d9d9d9")
        self.Welcome_Label.configure(disabledforeground="#a3a3a3")
        self.Welcome_Label.configure(font="-family {Arial} -size 20")
        self.Welcome_Label.configure(foreground="#000000")
        self.Welcome_Label.configure(highlightbackground="#d9d9d9")
        self.Welcome_Label.configure(highlightcolor="black")
        self.Welcome_Label.configure(text='''Welcome to Deversorium''')
        self.Welcome_Label.configure(justify='center')

        self.Main_Frame = tk.Frame(top)
        self.Main_Frame.place(relx=0.142, rely=0.254, relheight=0.734
                , relwidth=0.850)
        self.Main_Frame.configure(borderwidth="2")
        self.Main_Frame.configure(background="#d9d9d9")

        self.Home_Button = tk.Button(top)
        self.Home_Button.place(relx=0.015, rely=0.811, height=31, width=71)
        self.Home_Button.configure(activebackground="#ececec")
        self.Home_Button.configure(activeforeground="#000000")
        self.Home_Button.configure(background="#d9d9d9")
        self.Home_Button.configure(disabledforeground="#a3a3a3")
        self.Home_Button.configure(font=font14)
        self.Home_Button.configure(foreground="#000000")
        self.Home_Button.configure(highlightbackground="#d9d9d9")
        self.Home_Button.configure(highlightcolor="black")
        self.Home_Button.configure(pady="0")
        self.Home_Button.configure(text='''Home''')
        self.Home_Button.configure(command=lambda:self.welcome_message(message,font1))
        self.Home_Button.configure(cursor='hand2')

    def get_table(self,db_name,tables):
        '''This method will display a table with read only permissions'''
        clear_frame_or_window(self.Main_Frame)

        
        nb = ttk.Notebook(self.Main_Frame)
        ## Below is a for loop that will create a tab for each element in tables.
        for table in tables:
            frame=tk.Frame(nb,background="#d9d9d9")
            read_only = ReadOnlyPermission(frame,nb)
            read_only.create_treeview(db_name,table)
            nb.add(frame,text=create_space(table),underline=0)
        nb.pack(expand=1,fill="both")
        nb.enable_traversal()
        
            
    def welcome_message(self,message,font):
        '''This will display a message in the main frame'''
        clear_frame_or_window(self.Main_Frame)
        text = tk.Message(self.Main_Frame,text=message,background="#d9d9d9",anchor=tk.NW,aspect=1000,
                          font=font, foreground = "#000000")
        #text.configure(relief="groove")
        ## Message box options 'https://www.python-course.eu/tkinter_message_widget.php'.
        text.pack(expand=1,fill=tk.BOTH)
           
class ManagementMainScreen(MainMenu):
    def __init__(self,top=None, message=None,users_record=None):
        '''This class populates the manager's main screen'''

        super().__init__(top, message, users_record) # calling parent class. 

        font11 = "-family Arial -size 8 -weight normal -slant roman "  \
            "-underline 1 -overstrike 0"
        font12 = "-family Arial -size 9 -weight normal -slant roman "  \
            "-underline 1 -overstrike 0"
        font13 = "-family Arial -size 10 -weight normal -slant roman "  \
            "-underline 1 -overstrike 0"

        self.top = top

        ## private attribute. (destroy in all subclasses)
        self._Month_Graph_Button = tk.Button(top)
        self._Month_Graph_Button.place(relx=0.015, rely=0.444, height=31, width=71)
        self._Month_Graph_Button.configure(activebackground="#ececec")
        self._Month_Graph_Button.configure(activeforeground="#000000")
        self._Month_Graph_Button.configure(background="#d9d9d9")
        self._Month_Graph_Button.configure(disabledforeground="#a3a3a3")
        self._Month_Graph_Button.configure(font=font11)
        self._Month_Graph_Button.configure(foreground="#000000")
        self._Month_Graph_Button.configure(highlightbackground="#d9d9d9")
        self._Month_Graph_Button.configure(highlightcolor="#000000")
        self._Month_Graph_Button.configure(pady="0")
        self._Month_Graph_Button.configure(text='''Month graph''')
        self._Month_Graph_Button.configure(cursor='hand2')
        self._Month_Graph_Button.configure(command= lambda: self.month_graph(True))
        
        ## private attribute. (destroy in all subclasses)
        self._Room_Graph_Button = tk.Button(top)
        self._Room_Graph_Button.place(relx=0.015, rely=0.533, height=31, width=71)
        self._Room_Graph_Button.configure(activebackground="#ececec")
        self._Room_Graph_Button.configure(activeforeground="#000000")
        self._Room_Graph_Button.configure(background="#d9d9d9")
        self._Room_Graph_Button.configure(disabledforeground="#a3a3a3")
        self._Room_Graph_Button.configure(font=font11)
        self._Room_Graph_Button.configure(foreground="#000000")
        self._Room_Graph_Button.configure(highlightbackground="#d9d9d9")
        self._Room_Graph_Button.configure(highlightcolor="black")
        self._Room_Graph_Button.configure(pady="0")
        self._Room_Graph_Button.configure(text='''Room graph''')
        self._Room_Graph_Button.configure(cursor='hand2')
        self._Room_Graph_Button.configure(command= lambda: self.room_graph(True))

        self.Tape_Chart_Button = tk.Button(top)
        self.Tape_Chart_Button.place(relx=0.015, rely=0.356, height=31, width=71)
        self.Tape_Chart_Button.configure(activebackground="#ececec")
        self.Tape_Chart_Button.configure(activeforeground="#000000")
        self.Tape_Chart_Button.configure(background="#d9d9d9")
        self.Tape_Chart_Button.configure(disabledforeground="#a3a3a3")
        self.Tape_Chart_Button.configure(font=font12)
        self.Tape_Chart_Button.configure(foreground="#000000")
        self.Tape_Chart_Button.configure(highlightbackground="#d9d9d9")
        self.Tape_Chart_Button.configure(highlightcolor="black")
        self.Tape_Chart_Button.configure(pady="0")
        self.Tape_Chart_Button.configure(text='''Tape Chart''')
        self.Tape_Chart_Button.configure(cursor='hand2')
        self.Tape_Chart_Button.configure(command=self.place_gantt_chart)
        
        ## private attribute. (destroy in all subclasses)
        self._Database_Button = tk.Button(top)
        self._Database_Button.place(relx=0.015, rely=0.267, height=31, width=71)
        self._Database_Button.configure(activebackground="#ececec")
        self._Database_Button.configure(activeforeground="#000000")
        self._Database_Button.configure(background="#d9d9d9")
        self._Database_Button.configure(disabledforeground="#a3a3a3")
        self._Database_Button.configure(font=font13)
        self._Database_Button.configure(foreground="#000000")
        self._Database_Button.configure(highlightbackground="#d9d9d9")
        self._Database_Button.configure(highlightcolor="black")
        self._Database_Button.configure(pady="0")
        self._Database_Button.configure(text='''Database''')
        self._Database_Button.configure(command=lambda:self.get_database("Hotel_Management_System.db",retrn_table_names("Hotel_Management_System.db")))
        self._Database_Button.configure(cursor='hand2')
        
        ## private attribute. (destroy in all subclasses)
        self._Report_Button = tk.Button(top)
        self._Report_Button.place(relx=0.015, rely=0.622, height=31, width=71)
        self._Report_Button.configure(activebackground="#ececec")
        self._Report_Button.configure(activeforeground="#000000")
        self._Report_Button.configure(background="#d9d9d9")
        self._Report_Button.configure(disabledforeground="#a3a3a3")
        self._Report_Button.configure(font=font13)
        self._Report_Button.configure(foreground="#000000")
        self._Report_Button.configure(highlightbackground="#d9d9d9")
        self._Report_Button.configure(highlightcolor="black")
        self._Report_Button.configure(pady="0")
        self._Report_Button.configure(text='''Report''')
        self._Report_Button.configure(cursor='hand2')
        self._Report_Button.configure(command=self.create_report)

        ## private attribute. (destroy in all subclasses)
        self._Backup_Button = tk.Button(top)
        self._Backup_Button.place(relx=0.015, rely=0.722, height=31, width=71)
        self._Backup_Button.configure(activebackground="#ececec")
        self._Backup_Button.configure(activeforeground="#000000")
        self._Backup_Button.configure(background="#d9d9d9")
        self._Backup_Button.configure(disabledforeground="#a3a3a3")
        self._Backup_Button.configure(font=font11)
        self._Backup_Button.configure(foreground="#000000")
        self._Backup_Button.configure(highlightbackground="#d9d9d9")
        self._Backup_Button.configure(highlightcolor="#000000")
        self._Backup_Button.configure(pady="0")
        self._Backup_Button.configure(text='''Backup''')
        self._Backup_Button.configure(cursor='hand2')
        self._Backup_Button.configure(command=self.backup_pressed)

    def create_report(self):
        '''This will create the report at a specific directory and will display the
           updated graphs'''
        
        report = docx.Document()

        today = datetime.today() # stores todays date and time the script was run.
        current_date = today.strftime("%d-%m-%Y")

        header = report.sections[0].header
        header = header.add_paragraph('This report was created on the {0}'.format(current_date))
        header_font = header.runs[0].font
        header_font.name = 'Cambria (Body)'
        header_font.size = Pt(11)

        heading = report.add_heading("Hotel Report                                       ",0)
        pic = heading.add_run()
        pic.add_picture("hotel_logo.png",width=docx.shared.Inches(0.75))

        ## Below sets font for the whole document.
        style = report.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        info_on_reprot = report.add_paragraph("This reprot includes the 2 graphs that can be accessed within the main system.")
        #heading.style = report.styles['Normal']
        #heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #font_for_heading = heading.runs[0].font
        #font_for_heading.name = "Arial"
        #font_for_heading.size = Pt(14)
        #font_for_heading.bold = True

        message = report.add_heading("Month graph:",1)
        report.add_paragraph("The graph below shows the percentage of bookings per month.")
        pie_chart_name = self.month_graph(False)
        img = Image.open(pie_chart_name)
        img1 = img.crop((100, 50, 550, 440)) # left, up, right, bottom
        img1.save(pie_chart_name)
        pie_chart = report.add_picture(pie_chart_name,height = docx.shared.Inches(3))
        pie_chart.alignment = WD_ALIGN_PARAGRAPH.CENTER

        message = report.add_heading("Room graph:",1)
        report.add_paragraph("The graph below shows the number of bookings per room.")
        bar_chart_name = self.room_graph(False)
        img = Image.open(bar_chart_name)
        img1 = img.crop((0, 20, 640, 450)) # left, up, right, bottom
        img1.save(bar_chart_name)
        bar_chart = message.add_run()
        report.add_picture(bar_chart_name, height = docx.shared.Inches(2.7))
        bar_chart.alignment = WD_ALIGN_PARAGRAPH.CENTER

        #drctory_of_backup = filedialog.askdirectory(title = "Select the folder you would like the backup to save too.")
        #file = drctory_of_backup+"Report.docx"
        #file = "Report.docx"

        files = files = [('Word Document', '*.docx')]

        file = filedialog.asksaveasfile(filetypes = files, defaultextension = files)

        try:
                report.save(file.name)
        except AttributeError:
                messagebox.showerror("Did not save","You did not chose a name for the file and therefore will not save")

    def backup_pressed(self):
        '''This will display a message in the main frame of the system and display a backup
           button'''
        
        clear_frame_or_window(self.Main_Frame)
        message ="""This is the backup screen:
    -When any users logs onto the system the system will check when the last backup
     was and if it was over a month ago then the system back itself up
    -This backup is saved within the same folder as the system within a folder called 'backup'
    -When there are 3 backups withinthis folder and a 4th backup is made the latest backup will
     be deleted and therefore the total of backups within the folder will remain 3
        -The reason for this is to save storage space
    -When the backup button is pressed below you will need to:
        -Choose a FOLDER where you would like the backup to be saved
        -At the chosen location a folder called 'backup' will be there
        -The backup will be in this folder
    -The backup will have the same name as the original except with the date it was backed-up at
     the end"""

        font = "-family Arial -size 10 -weight normal -slant roman "  \
            "-underline 0 -overstrike 0"

        button_font = "-family Arial -size 20 -weight normal -slant roman "  \
            "-underline 1 -overstrike 0"
        
        self.welcome_message(message, font)
        backup_button=tk.Button(self.Main_Frame,text="Backup",command=get_file_location, font=button_font)
        backup_button.configure(background="#d9d9d9")
        backup_button.configure(width=10,height=2)
        backup_button.pack(pady=20)

    def get_database(self,db_name,tables):
        '''This will create a notebook with each tab being a table.
           Each table has been given read and write permission'''
        clear_frame_or_window(self.Main_Frame)
        nb = ttk.Notebook(self.Main_Frame)
        ## Below is a for loop that will create a tab for each element in tables.
        for table in tables:
            if table == "Room" and self.__class__.__name__ == "ReceptionistMainScreen":
                    frame=tk.Frame(nb,background="#d9d9d9")
                    read_only = ReadOnlyPermission(frame,nb)
                    read_only.create_treeview(db_name,table)
            else:
                    frame=tk.Frame(nb,background="#d9d9d9")
                    read_write = ReadAndWritePermission(frame,nb)
                    read_write.create_treeview(db_name,table)
            """Will only underline the first character and
               as some of the tables have the same first digit it will only
               go to the first one with that character (this is for underline).
               This has been taken from 'https://anzeljg.github.io/rin2/book2/2405/docs/tkinter/ttk-Notebook.html'"""
            if table == "Bill" or table == "Room_Facilities":
                    nb.add(frame,text=create_space(table),underline=1)
            elif table == "Room_Type":
                    nb.add(frame,text=create_space(table),underline=5)
            elif table == "Staff_Room":
                    nb.add(frame,text=create_space(table),underline=2)
            else:
                    nb.add(frame,text=create_space(table),underline=0)
            #nb.configure(cursor="hand2")
        nb.pack(expand=1,fill="both")
        nb.enable_traversal()
        return True

    def month_graph(self, display):
        '''This will get the relevant data from the database required to display
           the month graph, when display is false it will only return the image name'''
        labels = []
        
        sizes = []

        months_and_counter = [["Jan",0],["Feb",0],["Mar",0],["Apr",0],["May",0],["Jun",0],["Jul",0],["Aug",0],["Sep",0],["Oct",0],["Nov",0],["Dec",0]]
        with sqlite3.connect("Hotel_management_system.db") as db:
            cursor = db.cursor()
            booking_sql = cursor.execute("select Check_in from Booking")
            bookings = cursor.fetchall()
            
        for i in bookings:
            check_in = i[0]
            for i in range(len(check_in)):
                if check_in[i] == "/":
                    month = int(check_in[i+1])
                    if month == 0 or (month == 1 and check_in[i+2] != "/"):
                            month = int(check_in[i+1]+check_in[i+2])
                    months_and_counter[month-1][1]=months_and_counter[month-1][1]+1
                    break
        for i in months_and_counter:
            if i[1] != 0:
                labels.append(i[0])
                sizes.append(i[1])

        img_name = pie_chart(labels,sizes)

        if display == False:
            return img_name
        else:
            set_image(self.Main_Frame, img_name,600, 400)

    def room_graph(self,display):
        '''This will get the relevant data from the database required to display
           the room graph, when display is false it will only return the image name'''
        yaxis = []
        xaxis = []

        counter = [[0,0]]
        
        with sqlite3.connect("Hotel_management_system.db") as db:
            cursor = db.cursor()
            booking_sql = cursor.execute("select RoomID from Booking")
            bookings = cursor.fetchall()
            for i in bookings:
                rooms_sql = cursor.execute("select Room_Number from Room where RoomID = {0}".format(i[0]))
                room_number = cursor.fetchone()
                for i in counter:
                    if room_number[0] == i[0]:
                        i[1] = i[1]+1
                        break
                    elif len(counter) == counter.index(i)+1:
                        counter.append([room_number[0],1])
                        break

            counter.remove([0,0])
            
            for i in counter:
                if i[1] != 0:
                    xaxis.append("Room {0}".format(i[0]))
                    yaxis.append(i[1])
            xaxis.sort()

            img_name = bar_chart(xaxis, yaxis)

        if display == False:
            return img_name
        else:
            set_image(self.Main_Frame, img_name,600, 350)

    def place_gantt_chart(self):
        '''This places the drop down menu, the enrty widget and the button required
           to create the 'Tape Chart' '''
        
        clear_frame_or_window(self.Main_Frame)

        frame_buttons = tk.Frame(self.Main_Frame)
        frame_buttons.configure(background="#d9d9d9")
        frame_buttons.pack(side = tk.TOP)

        frame_image = tk.Frame(self.Main_Frame)
        frame_image.configure(background="#d9d9d9")
        frame_image.pack(side = tk.TOP)

        year = tk.Entry(frame_buttons)
        year.grid(row=0,column=1)
        year_label = tk.Label(frame_buttons,text = "Year:",background="#d9d9d9")
        year_label.grid(row=0,column=0)

        today = datetime.today()

        year.insert(0,str(today.year))

        var = tk.StringVar(frame_buttons)
        all_months = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
        var.set(all_months[today.month-1])

        month_menu = tk.OptionMenu(frame_buttons,var, *all_months)
        month_menu.configure(background="#d9d9d9")
        month_menu.configure(cursor="hand2")
        month_menu.grid(row=1,column=1,sticky="ew")
        month_label = tk.Label(frame_buttons,text = "Month:",background="#d9d9d9")
        month_label.grid(row=1,column=0,padx=5)

    
        enter_button = tk.Button(frame_buttons, text = "Go",command = lambda: set_image(frame_image,get_data_for_gantt(all_months.index(var.get())+1,year.get()),500,300))
        enter_button.configure(background="#d9d9d9")
        enter_button.grid(row=1,column=2,padx=5)

        try:
            self.top.bind('<Return>', lambda key_pressed: set_image(frame_image,get_data_for_gantt(all_months.index(var.get())+1,year.get()),500,300))
        except _tkinter.TclError:
            pass
        
        set_image(frame_image,get_data_for_gantt(today.month,today.year),500,290)

        
class CustomerMainScreen(MainMenu):
    def __init__(self, top=None, message=None, users_record=None):
        '''This populates the customer's main screen'''
        
        super().__init__(top, message, users_record) # calling parent class.

        font14 = "-family Arial -size 10 -weight bold -slant roman "  \
                 "-underline 1 -overstrike 0"
        font15 = "-family Arial -size 7 -weight normal -slant roman "  \
                 "-underline 1 -overstrike 0"
        font16 = "-family Arial -size 6 -weight normal -slant roman "  \
                 "-underline 1 -overstrike 0"

        ## private attribute. (destroy in all subclasses)
        self._Check_Booking_Button = tk.Button(top)
        self._Check_Booking_Button.place(relx=0.015, rely=0.356, height=31, width=71)
        self._Check_Booking_Button.configure(activebackground="#ececec")
        self._Check_Booking_Button.configure(activeforeground="#000000")
        self._Check_Booking_Button.configure(background="#d9d9d9")
        self._Check_Booking_Button.configure(disabledforeground="#a3a3a3")
        self._Check_Booking_Button.configure(font=font15)
        self._Check_Booking_Button.configure(foreground="#000000")
        self._Check_Booking_Button.configure(highlightbackground="#d9d9d9")
        self._Check_Booking_Button.configure(highlightcolor="black")
        self._Check_Booking_Button.configure(pady="0")
        self._Check_Booking_Button.configure(text='''Check Booking''')
        self._Check_Booking_Button.configure(cursor='hand2')
        self._Check_Booking_Button.configure(command=lambda:self.request_record("Hotel_Management_System.db",users_record[0]))

        ## private attribute. (destroy in all subclasses)
        self._Book_Button = tk.Button(top)
        self._Book_Button.place(relx=0.015, rely=0.267, height=31, width=71)
        self._Book_Button.configure(activebackground="#ececec")
        self._Book_Button.configure(activeforeground="#000000")
        self._Book_Button.configure(background="#d9d9d9")
        self._Book_Button.configure(disabledforeground="#a3a3a3")
        self._Book_Button.configure(font="-family {Arial} -size 9 -underline 1")
        self._Book_Button.configure(foreground="#000000")
        self._Book_Button.configure(highlightbackground="#d9d9d9")
        self._Book_Button.configure(highlightcolor="black")
        self._Book_Button.configure(pady="0")
        self._Book_Button.configure(text='''Book''')
        self._Book_Button.configure(command=lambda:self.book("Hotel_Management_System.db","Room",users_record))
        self._Book_Button.configure(cursor='hand2')

        self.Request_Record_Button = tk.Button(top)
        self.Request_Record_Button.place(relx=0.015, rely=0.444, height=31, width=71)
        self.Request_Record_Button.configure(activebackground="#ececec")
        self.Request_Record_Button.configure(activeforeground="#000000")
        self.Request_Record_Button.configure(background="#d9d9d9")
        self.Request_Record_Button.configure(disabledforeground="#a3a3a3")
        self.Request_Record_Button.configure(font=font16)
        self.Request_Record_Button.configure(foreground="#000000")
        self.Request_Record_Button.configure(highlightbackground="#d9d9d9")
        self.Request_Record_Button.configure(highlightcolor="black")
        self.Request_Record_Button.configure(pady="0")
        self.Request_Record_Button.configure(text='''Request Record''')
        self.Request_Record_Button.configure(cursor='hand2')
        self.Request_Record_Button.configure(command=lambda:self.request_record("Hotel_Management_System.db",users_record))

    def book(self,db_name,table,users_record):
        '''This will only display the room table with a button
           that allows the user to book a holiday'''
        
        clear_frame_or_window(self.Main_Frame)
        nb = ttk.Notebook(self.Main_Frame)
        frame=tk.Frame(nb,background="#d9d9d9")
        book_hol = BookHoliday(frame,nb,table,users_record)
        book_hol.create_treeview(db_name,table)
        nb.add(frame,text=table,underline=0)
        nb.pack(expand=1,fill="both")
        nb.enable_traversal()

    def request_record(self,db_name,users_record):
        '''This will display a notebook containing a table with only the
           user's record'''
        
        if self.__class__.__name__ == "CustomerMainScreen":
                table = "Customer"
        else:
                table = "Staff"
        clear_frame_or_window(self.Main_Frame)
        nb = ttk.Notebook(self.Main_Frame)
        frame=tk.Frame(nb,background="#d9d9d9")
        if type(users_record) == tuple:
                nb.add(frame,text=table,underline=0)
                request_record = RequestRecord(frame,nb,table,users_record)
                request_record.record_tview(db_name,table)
        else:
                with sqlite3.connect(db_name) as db:
                        cursor=db.cursor()
                        cursor.execute("select * from Booking where CustomerID={0}".format(users_record))
                        booking_record=cursor.fetchall() # gets all records form table.
                        db.commit()
                        
                for booking in booking_record:
                        today = datetime.today() # stores todays date and time the script was run.
                        check_in = datetime.strptime(booking[2],"%d/%m/%Y")
                        if today>check_in:
                                booking_record.remove(booking)
                                
                nb.add(frame,text="Booking",underline=0)
                request_record = RequestRecord(frame,nb,"Booking",booking_record)
                request_record.record_tview(db_name,"Booking")
        nb.pack(expand=True,fill=tk.BOTH)
            

class CleanerMainScreen(CustomerMainScreen):
    def __init__(self, top=None, message=None, users_record=None):
        '''This populates the cleaner's main screen.'''
        
        super().__init__(top, message, users_record) # calling parent class.

        ## destroying all private 
        self._Check_Booking_Button.destroy()
        self._Book_Button.destroy()

        self.Room_Table_Button = tk.Button(top)
        self.Room_Table_Button.place(relx=0.015, rely=0.267, height=31, width=71)
        self.Room_Table_Button.configure(activebackground="#ececec")
        self.Room_Table_Button.configure(activeforeground="#000000")
        self.Room_Table_Button.configure(background="#d9d9d9")
        self.Room_Table_Button.configure(disabledforeground="#a3a3a3")
        self.Room_Table_Button.configure(font="-family {Arial} -size 9 -underline 1")
        self.Room_Table_Button.configure(foreground="#000000")
        self.Room_Table_Button.configure(highlightbackground="#d9d9d9")
        self.Room_Table_Button.configure(highlightcolor="black")
        self.Room_Table_Button.configure(pady="0")
        self.Room_Table_Button.configure(text='''Room Table''')
        self.Room_Table_Button.configure(command=lambda:self.get_table("Hotel_Management_System.db",["Room"]))
        self.Room_Table_Button.configure(cursor='hand2')

class AccountantsMainScreen(CustomerMainScreen):
    def __init__(self, top=None, message=None, users_record=None):
        '''This populates the accountant's main screen.'''
        
        super().__init__(top, message, users_record) # calling parent class.

        ## destroying all private 
        self._Check_Booking_Button.destroy()
        self._Book_Button.destroy()
        
        font13 = "-family Arial -size 10 -weight normal -slant roman "  \
                 "-underline 1 -overstrike 0"
        font14 = "-family Arial -size 10 -weight bold -slant roman "  \
                 "-underline 1 -overstrike 0"
        font15 = "-family Arial -size 7 -weight normal -slant roman "  \
                 "-underline 1 -overstrike 0"
        
        self.Bill_Table_Button = tk.Button(top)
        self.Bill_Table_Button.place(relx=0.015, rely=0.267, height=31, width=71)
        self.Bill_Table_Button.configure(activebackground="#ececec")
        self.Bill_Table_Button.configure(activeforeground="#000000")
        self.Bill_Table_Button.configure(background="#d9d9d9")
        self.Bill_Table_Button.configure(disabledforeground="#a3a3a3")
        self.Bill_Table_Button.configure(font="-family {Arial} -size 9 -underline 1")
        self.Bill_Table_Button.configure(foreground="#000000")
        self.Bill_Table_Button.configure(highlightbackground="#d9d9d9")
        self.Bill_Table_Button.configure(highlightcolor="black")
        self.Bill_Table_Button.configure(pady="0")
        self.Bill_Table_Button.configure(text='''Bill Table''')
        self.Bill_Table_Button.configure(command=lambda:self.get_table("Hotel_Management_System.db",["Bill"]))
        self.Bill_Table_Button.configure(cursor='hand2')

        self.Salary_Button = tk.Button(top)
        self.Salary_Button.place(relx=0.015, rely=0.356, height=31, width=71)
        self.Salary_Button.configure(activebackground="#ececec")
        self.Salary_Button.configure(activeforeground="#000000")
        self.Salary_Button.configure(background="#d9d9d9")
        self.Salary_Button.configure(disabledforeground="#a3a3a3")
        self.Salary_Button.configure(font=font13)
        self.Salary_Button.configure(foreground="#000000")
        self.Salary_Button.configure(highlightbackground="#d9d9d9")
        self.Salary_Button.configure(highlightcolor="black")
        self.Salary_Button.configure(pady="0")
        self.Salary_Button.configure(text='''Salary''')
        self.Salary_Button.configure(cursor='hand2')
        self.Salary_Button.configure(command = self.get_salaries)

    def get_salaries(self):
        '''This will get all the salaries of the staff and create the average salary for that position,
           this data will then be displayed in the system'''
        all_salaries = []
        positions = ["Manager","Cleaner","Accountant","Receptionist"]
        for position in positions:
            all_salaries.append(get_av_salary(position))

        text = """
Below is the average salary for each position:
    -Manager = {0}
    -Cleaner = {1}
    -Accountant = {2}
    -Receptionist = {3}

        """.format(all_salaries[0],all_salaries[1],all_salaries[2],all_salaries[3])
        
        font = "-family Arial -size 14 -weight normal -slant roman "  \
            "-underline 0 -overstrike 0"
        self.welcome_message(text, font)

    
class ReceptionistMainScreen(CustomerMainScreen,ManagementMainScreen):
    def __init__(self, top=None, message=None, users_record=None):
        '''This populates Receptionist's main screen'''
        
        super().__init__(top, message, users_record) # calling parent class.

        ## destroying all private 
        self._Check_Booking_Button.destroy()
        self._Book_Button.destroy()
        self._Month_Graph_Button.destroy()
        self._Room_Graph_Button.destroy()
        self._Database_Button.destroy()
        self._Report_Button.destroy()
        self._Backup_Button.destroy()

        font11 = "-family Arial -size 8 -weight normal -slant roman "  \
                 "-underline 1 -overstrike 0"

        self.Check_Tables_Button = tk.Button(top)
        self.Check_Tables_Button.place(relx=0.015, rely=0.267, height=31, width=71)
        self.Check_Tables_Button.configure(activebackground="#ececec")
        self.Check_Tables_Button.configure(activeforeground="#000000")
        self.Check_Tables_Button.configure(background="#d9d9d9")
        self.Check_Tables_Button.configure(disabledforeground="#a3a3a3")
        self.Check_Tables_Button.configure(font=font11)
        self.Check_Tables_Button.configure(foreground="#000000")
        self.Check_Tables_Button.configure(highlightbackground="#d9d9d9")
        self.Check_Tables_Button.configure(highlightcolor="black")
        self.Check_Tables_Button.configure(pady="0")
        self.Check_Tables_Button.configure(text='''Check Tables''')
        self.Check_Tables_Button.configure(command=lambda:self.get_database("Hotel_Management_System.db",["Room","Customer","Booking"]))
        self.Check_Tables_Button.configure(cursor='hand2')


            
if __name__ == '__main__':
    ## Below is code for testing purposes.
    ## This code allowed me to bypass my login screen and go straight to the manager's screen
    
    root = tk.Tk()

    message ="""This is the main screen of the hotel Management system:
-To quit simply select the quit button at the bottom left.
-If you wish to return to this page at any time simply select the 'Home' button at the bottom left.
-To view the database in its entirety please hit the 'Database' button to the left. This will display the database with
 each tab being a table in the database.
    -In this screen you can add,edit,delete and search for each record
        -To delete a record just select the record and hit the edit button
        -To edit a record just select the record and hit edit
                -A pop-up window should appear and simply type the changes you would like to make
        -To add a record hit the add button
                -A pop-up window should appear and simply type the record you wish to add
        -To search for a record select the search button
                -A pop-ip window will appear and will prompt you for the 'ID' of the record you would like to search
        -There is an advanced search option allows you to search by any field
                -A pop-ip window will appear and will prompt you to choose a field and to enter in a search term.
                -A new tab will open and with the search results. 
-When the tape chart button is selected this will show a chart of when the rooms are occupied.
-The 'Month graph' button will display a graph of number of bookings per months.
    -This is based of how many check-ins there are per month.
-The 'Room graph' button will display the number of guests per room.
-The 'Report' button will produce a reprot which contian:
    -The 2 graphs described above.
    -The date which the report was created. 
-More information on backups can be found when the 'Backup' button is pressed."""
    mang = ManagementMainScreen(root,message, None)

    font = "-family Arial -size 8 -weight normal -slant roman "  \
            "-underline 0 -overstrike 0"
    
    mang.welcome_message(message,font)
    
    root.mainloop()
